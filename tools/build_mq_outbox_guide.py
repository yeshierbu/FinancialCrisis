from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from pathlib import Path


OUT = Path("artifacts/FinancialCrisis-MQ-Outbox链路与重试机制代码导读.docx")
BLUE = "1F4E78"
MID_BLUE = "D9EAF7"
LIGHT_BLUE = "EDF4FA"
LIGHT_GRAY = "F2F4F7"
GRAY = "666666"
GREEN = "E2F0D9"
YELLOW = "FFF2CC"
RED = "FCE4D6"
WHITE = "FFFFFF"
INK = "243447"
FONT = "Noto Sans SC"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_widths(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    total = sum(widths)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[min(idx, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_run(run, size=11, bold=False, color=INK, italic=False, font=FONT):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_text(doc, text, bold_prefix=None, after=6, keep=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.keep_with_next = keep
    if bold_prefix and text.startswith(bold_prefix):
        set_run(p.add_run(bold_prefix), bold=True)
        set_run(p.add_run(text[len(bold_prefix):]))
    else:
        set_run(p.add_run(text))
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.2
    set_run(p.add_run(text))
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.2
    set_run(p.add_run(text))
    return p


def add_callout(doc, label, text, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_widths(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.2
    set_run(p.add_run(label + "  "), bold=True, color=BLUE)
    set_run(p.add_run(text))
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(doc, headers, rows, widths, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, value in enumerate(headers):
        set_cell_shading(hdr.cells[i], MID_BLUE)
        hdr.cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = hdr.cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(value), size=font_size, bold=True, color=BLUE)
    for row_idx, values in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_idx % 2:
                set_cell_shading(cells[i], "F8FAFC")
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.12
            if i == 0 and len(headers) > 2:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run(p.add_run(str(value)), size=font_size)
    set_table_widths(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_code(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_widths(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F6F8FA")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    for idx, line in enumerate(lines.splitlines()):
        run = p.add_run(line)
        set_run(run, size=8.5, font=FONT, color="2F3B45")
        if idx < len(lines.splitlines()) - 1:
            run.add_break()
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_flow(doc, rows):
    table = doc.add_table(rows=0, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, (text, fill) in enumerate(rows):
        cell = table.add_row().cells[0]
        set_cell_shading(cell, fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(text), size=10, bold=True if idx % 2 == 0 else False,
                color=BLUE if idx % 2 == 0 else GRAY)
    set_table_widths(table, [7600])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    set_run(p.add_run(text), size={1: 16, 2: 13, 3: 11.5}[level], bold=True,
            color=BLUE if level < 3 else "365F91")
    return p


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(paragraph.add_run("第 "), size=9, color=GRAY)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)
    set_run(paragraph.add_run(" 页"), size=9, color=GRAY)


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, before, after in (("Heading 1", 16, 18, 10), ("Heading 2", 13, 14, 7), ("Heading 3", 11.5, 10, 5)):
        st = styles[name]
        st.font.name = FONT
        st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(BLUE)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(header.add_run("FinancialCrisis · MQ 技术导读"), size=9, color=GRAY)
    add_page_number(section.footer.paragraphs[0])

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(42)
    p.paragraph_format.space_after = Pt(8)
    set_run(p.add_run("FinancialCrisis 项目"), size=13, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    set_run(p.add_run("MQ Outbox 链路与重试机制"), size=26, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(24)
    set_run(p.add_run("结合源码的实现说明、异常分支与调试指南"), size=14, color=GRAY)
    add_callout(doc, "核心结论", "OCR 结果写入业务表；材料齐全后写 approval_outbox。Outbox 发布失败由数据库扫描重试，审批消费失败由 RabbitMQ 延迟重试；正常业务异常达到上限后转人工并进入最终死信队列。", MID_BLUE)
    add_text(doc, "适用对象：希望理解本项目 RabbitMQ、Transactional Outbox、手动 ACK、幂等消费及死信机制的开发者。", after=3)
    add_text(doc, "代码基线：backend/src/main/java 与 backend/src/main/resources/mapper。", after=3)
    add_text(doc, "整理日期：2026-07-21", after=3)

    doc.add_page_break()
    add_heading(doc, "1. 先建立正确的链路认知", 1)
    add_callout(doc, "一句话", "扫描的是 approval_outbox，不是 OCR 材料表；消费失败不会把原 Outbox 改回 PENDING，而是产生 RabbitMQ 延迟重试消息。", YELLOW)
    add_flow(doc, [
        ("用户上传材料", MID_BLUE),
        ("↓", WHITE),
        ("OCR 识别并保存 uploaded_document", LIGHT_BLUE),
        ("↓ 材料齐全", WHITE),
        ("同一 MySQL 事务：申请状态 → SUBMITTED + Outbox → PENDING", GREEN),
        ("↓ 定时扫描 approval_outbox", WHITE),
        ("Publisher → approval.exchange → approval.start.queue", MID_BLUE),
        ("↓", WHITE),
        ("Consumer 每次执行一个 ApprovalStep", LIGHT_BLUE),
        ("↓ 成功：下一步骤 Outbox / 失败：RabbitMQ Retry", WHITE),
        ("完成审批，或重试耗尽后转人工 + DLQ", GREEN),
    ])
    add_heading(doc, "1.1 两类数据库数据", 2)
    add_table(doc, ["数据类别", "主要内容", "作用", "是否被发布器扫描"], [
        ["OCR 业务数据", "材料类型、OCR 状态、识别结果 JSON、文件哈希", "供审批步骤按 applicationId 查询", "否"],
        ["Outbox 事件", "eventId、applicationId、step、routingKey、payloadJson", "代表一件待发布的审批任务", "是"],
    ], [1700, 2900, 2860, 1900])
    add_heading(doc, "1.2 两类重试不能混淆", 2)
    add_table(doc, ["失败阶段", "恢复载体", "状态/队列", "含义"], [
        ["消息发布失败", "MySQL Outbox", "PUBLISHING → RETRY", "消息尚未可靠交给 Broker"],
        ["审批消费失败", "RabbitMQ 重试队列", "retry.queue → TTL → 主队列", "Broker 已有消息，但业务步骤执行失败"],
    ], [1900, 1900, 2500, 3060])

    add_heading(doc, "2. 第一段：OCR 与 Outbox 创建", 1)
    add_heading(doc, "2.1 业务入口", 2)
    add_text(doc, "上传入口位于 DocumentController.uploadDocument()；批量补充入口位于 DocumentController.submitSupplement()。两者最终进入 DocumentServiceImpl。")
    add_code(doc, "// DocumentController.java:34-45\nPOST /api/loan/applications/{applicationId}/documents\nPOST /api/loan/applications/{applicationId}/supplement")
    add_heading(doc, "2.2 OCR 结果先持久化", 2)
    add_text(doc, "DocumentServiceImpl.uploadDocument() 先保存材料元数据，再调用 OCR，并把识别结果和 SUCCESS/FAILED 状态更新到数据库。")
    add_code(doc, "// DocumentServiceImpl.java:96-102\nstore.addDocument(document);\ndocument.setParseResultJson(qianfanOcrService.recognize(documentType, dataUrl));\ndocument.setOcrStatus(OcrStatus.SUCCESS);\nstore.updateDocument(document);")
    add_text(doc, "若必需材料尚未全部成功，申请保持 DOCUMENT_PENDING，不创建审批任务；全部就绪后，申请改为 SUBMITTED，并调用 approvalTaskService.submit(applicationId)。")
    add_code(doc, "// DocumentServiceImpl.java:141-151\nstore.changeStatus(application, ApplicationStatus.SUBMITTED, ...);\napprovalTaskService.submit(applicationId);")
    add_heading(doc, "2.3 同一事务写 Outbox", 2)
    add_text(doc, "uploadDocument() 和 submitSupplement() 都标注 @Transactional。ApprovalTaskServiceImpl 没有开启新事务，因此 Outbox INSERT 会加入调用方当前事务。")
    add_callout(doc, "事务语义", "材料、OCR、申请状态、Outbox 要么一起提交，要么一起回滚。这解决了“数据库已提交但 MQ 消息没发出去”的双写不一致。", GREEN)
    add_code(doc, "// ApprovalTaskServiceImpl.java:65-81\nApprovalStartMessage message = new ApprovalStartMessage(\n    eventId, applicationId, application.getApplicationNo(), step, 0, now);\noutbox.setEventType(\"APPROVAL_STEP_\" + step.name());\noutbox.setRoutingKey(ApprovalRabbitConfig.APPROVAL_ROUTING_KEY);\noutbox.setPayloadJson(objectMapper.writeValueAsString(message));\noutbox.setPublishStatus(\"PENDING\");\noutboxMapper.insert(outbox);")

    add_heading(doc, "3. 第二段：Outbox 发布到主审批队列", 1)
    add_heading(doc, "3.1 Outbox 状态机", 2)
    add_flow(doc, [
        ("PENDING / RETRY", MID_BLUE), ("↓ claim 条件更新", WHITE),
        ("PUBLISHING", YELLOW), ("↓ Broker Confirm ACK", WHITE),
        ("PUBLISHED", GREEN), ("↓ 消费业务完成", WHITE),
        ("CONSUMED", GREEN),
    ])
    add_text(doc, "ApprovalOutboxPublisher.publishPending() 默认每 1 秒运行。它先恢复超过 1 分钟仍为 PUBLISHING 的记录，再选择到期的 PENDING/RETRY 事件。")
    add_code(doc, "// ApprovalOutboxPublisher.java:39-59\nmapper.resetStalePublishing(LocalDateTime.now().minusMinutes(1));\nfor (ApprovalOutbox event : mapper.selectPending(batchSize)) {\n    if (mapper.claim(event.getId()) != 1) continue;\n    rabbitTemplate.convertAndSend(APPROVAL_EXCHANGE, event.getRoutingKey(), payload, ...);\n    Confirm confirm = correlation.getFuture().get(10, TimeUnit.SECONDS);\n    if (!confirm.isAck()) throw new IllegalStateException(...);\n    mapper.markPublished(event.getId());\n}")
    add_heading(doc, "3.2 claim 为什么重要", 2)
    add_text(doc, "多个应用实例可能同时查到同一条 PENDING 记录。claim 使用带状态条件的 UPDATE，只有一个实例能把它改成 PUBLISHING；其他实例影响行数为 0，直接跳过。")
    add_code(doc, "-- ApprovalOutboxMapper.xml:17-20\nUPDATE approval_outbox\nSET publish_status='PUBLISHING', updated_at=CURRENT_TIMESTAMP\nWHERE id=#{id} AND publish_status IN ('PENDING', 'RETRY');")
    add_heading(doc, "3.3 发布失败如何恢复", 2)
    add_text(doc, "连接失败、Confirm NACK、Confirm 超时或序列化异常都会进入 catch，Outbox 改为 RETRY，并设置 next_retry_at。后续扫描器再次发送。")
    add_callout(doc, "注意", "Outbox 负责的是“可靠发布”，而不是消费者业务重试。只有消息还没可靠交给 RabbitMQ 时，才走 Outbox RETRY。", YELLOW)

    add_heading(doc, "4. 第三段：RabbitMQ 拓扑", 1)
    add_table(doc, ["用途", "Exchange", "Routing Key", "Queue", "行为"], [
        ["正常审批", "approval.exchange", "approval.start", "approval.start.queue", "消费者执行当前步骤"],
        ["延迟重试", "approval.retry.exchange", "approval.retry", "approval.retry.queue", "等待 TTL 后回投主交换机"],
        ["最终失败", "approval.dlx", "approval.dead", "approval.dead.queue", "保留消息，等待人工排查/补偿"],
    ], [1250, 2020, 1500, 2150, 2440], font_size=8.8)
    add_code(doc, "// ApprovalRabbitConfig.java:95-100\nQueueBuilder.durable(RETRY_QUEUE)\n    .ttl(delayMs)\n    .deadLetterExchange(APPROVAL_EXCHANGE)\n    .deadLetterRoutingKey(APPROVAL_ROUTING_KEY)\n    .build();")
    add_text(doc, "重试队列本身没有审批消费者。消息只是停留到 TTL 到期，然后 RabbitMQ 把它当作死信路由回 approval.exchange。这里的“死信转发”是延迟实现手段，不等于最终失败 DLQ。")

    add_heading(doc, "5. 第四段：主队列消费者", 1)
    add_text(doc, "ApprovalMessageConsumer.consume() 使用手动 ACK，按以下次序执行。")
    add_number(doc, "按 eventId 抢占消费权，拦截重复消息并提供租约接管。")
    add_number(doc, "按 applicationId 获取执行锁，防止同一申请的不同步骤并发运行。")
    add_number(doc, "检查申请是否已经 APPROVED、REJECTED、MANUAL_REVIEW 或 ARCHIVED。")
    add_number(doc, "执行 payload.step 对应的一个审批步骤。")
    add_number(doc, "存在 nextStep 时写入下一步骤 Outbox。")
    add_number(doc, "当前消费记录改为 COMPLETED、Outbox 改为 CONSUMED，再 ACK。")
    add_code(doc, "// ApprovalMessageConsumer.java:89-114\nString claimToken = consumeService.tryStart(payload, consumerName);\nif (claimToken == null) { basicAck(...); return; }\nif (!acquireApplicationLock(payload)) { ... }\nLoanApplication application = store.getApplicationOrThrow(payload.applicationId());\nvar nextStep = orchestrationService.executeStep(payload.applicationId(), payload.step());\nif (nextStep != null) taskService.submitStep(payload.applicationId(), nextStep, payload.eventId());\nconsumeService.complete(payload.eventId(), claimToken);\nchannel.basicAck(tag, false);")
    add_heading(doc, "5.1 三层防重/并发保护", 2)
    add_table(doc, ["机制", "键", "解决的问题", "关键实现"], [
        ["消费幂等", "eventId", "同一消息重复投递", "consume_log.event_id UNIQUE"],
        ["消费租约", "eventId + claimToken", "消费者宕机后的安全接管", "lease_until 过期可 retryClaim"],
        ["申请锁", "applicationId", "同申请的不同步骤并发", "execution_lock.application_id PK"],
        ["终态检查", "application.status", "迟到消息修改已结束申请", "终态直接完成并 ACK"],
    ], [1750, 1900, 2850, 2860])
    add_heading(doc, "5.2 一条消息只执行一个步骤", 2)
    add_text(doc, "ApprovalSingleStepExecutor.execute() 根据 step 分发。步骤成功时返回 nextStep；消费者将其写成新的 Outbox，而不是在当前消息里继续执行整条流程。")
    add_flow(doc, [
        ("DOCUMENT_INTAKE → DOCUMENT_ANALYSIS", MID_BLUE),
        ("↓", WHITE),
        ("FRAUD_ASSESSMENT → REPAYMENT_ASSESSMENT", LIGHT_BLUE),
        ("↓", WHITE),
        ("RISK_ANALYSIS → INDEPENDENT_REVIEW", MID_BLUE),
        ("↓ 不通过时仅返工一次", WHITE),
        ("RISK_REVISION → REVISION_REVIEW", YELLOW),
        ("↓", WHITE),
        ("FINAL_DECISION → POLICY_GUARD → 终态", GREEN),
    ])

    add_heading(doc, "6. 第五段：消费失败、重试与死信", 1)
    add_heading(doc, "6.1 正常业务异常：有限重试", 2)
    add_flow(doc, [
        ("主队列执行步骤失败", RED),
        ("↓ consume log → FAILED", WHITE),
        ("retryCount < maxRetries？", YELLOW),
        ("是：发送 nextRetry 到 retry.exchange，并 ACK 原消息", LIGHT_BLUE),
        ("↓ retry.queue 等待 30 秒，TTL 到期回主队列", WHITE),
        ("否：申请 → MANUAL_REVIEW + 创建工单 + 消息进入 DLQ", GREEN),
    ])
    add_code(doc, "// ApprovalMessageConsumer.java:141-150\nif (payload.retryCount() < maxRetries) {\n    rabbitTemplate.convertAndSend(RETRY_EXCHANGE, RETRY_ROUTING_KEY, payload.nextRetry());\n} else {\n    orchestrationService.moveToManualReview(payload.applicationId(), ...);\n    rabbitTemplate.convertAndSend(DEAD_EXCHANGE, DEAD_ROUTING_KEY, payload);\n}")
    add_text(doc, "默认 maxRetries=3。初始 retryCount=0，因此最多是首次执行加 3 次重试，共 4 次业务尝试。")
    add_table(doc, ["消费次数", "retryCount", "失败后的动作"], [
        ["第 1 次", "0", "发送 retryCount=1"],
        ["第 2 次", "1", "发送 retryCount=2"],
        ["第 3 次", "2", "发送 retryCount=3"],
        ["第 4 次", "3", "转人工审核，并发送到 approval.dead.queue"],
    ], [1900, 1900, 5560])
    add_heading(doc, "6.2 为什么发送重试成功后要 ACK 原消息", 2)
    add_text(doc, "因为系统已经创建了一个新的延迟重试消息。如果原消息不 ACK，它会立即重新投递，不仅绕过 30 秒延迟，还会与新重试消息形成重复执行。")
    add_heading(doc, "6.3 重试消息发送也失败", 2)
    add_code(doc, "// ApprovalMessageConsumer.java:121-123\ncatch (Exception publishFailure) {\n    channel.basicNack(tag, false, true);\n}")
    add_text(doc, "这时新 retryCount 消息没有成功发出，消费者对主队列原消息执行 NACK + requeue=true。原消息立即或稍后重新进入主队列，而且 retryCount 不增加。")
    add_callout(doc, "循环风险", "该分支没有次数上限和退避。如果重试/死信消息持续无法发布，原消息可能持续 NACK 回队，形成热循环。", RED)
    add_heading(doc, "6.4 申请锁竞争", 2)
    add_text(doc, "未获取 applicationId 锁时，代码认为这是临时并发冲突：标记当前消费失败，发送 nextRetry 到重试队列，再 ACK 原消息。这个分支没有 maxRetries 检查，因此持续锁冲突可能让 retryCount 一直增长。")
    add_callout(doc, "设计意图", "锁竞争不等于业务失败，不应轻易把正常申请转人工；但生产方案应增加最长等待时长、独立冲突计数或告警。", YELLOW)

    add_heading(doc, "7. 异常场景决策表", 1)
    add_table(doc, ["场景", "处理动作", "是否回主队列", "是否有限", "最终结果"], [
        ["Outbox 发布失败", "Outbox → RETRY，扫描器稍后重发", "重新发布后进入", "当前无最大次数", "Broker 恢复后继续"],
        ["审批业务失败，未达上限", "新消息进入 retry.queue，ACK 原消息", "TTL 到期后回主队列", "是，默认 3 次重试", "成功或继续重试"],
        ["审批业务失败，达到上限", "转人工并显式发送 DLQ", "否", "是", "MANUAL_REVIEW"],
        ["重试消息发送失败", "原消息 NACK + requeue", "是，原消息回队", "否", "恢复前可能循环"],
        ["死信消息发送失败", "原消息 NACK + requeue", "是，原消息回队", "否", "恢复前可能循环"],
        ["申请锁竞争", "nextRetry → retry.queue，ACK 原消息", "TTL 到期后回队", "代码无上限", "锁释放后继续"],
        ["消息进入最终 DLQ", "等待排查或补偿", "不会自动回主队列", "不适用", "人工处理"],
    ], [1850, 2950, 1780, 1400, 1380], font_size=8.3)
    add_heading(doc, "8. 为什么整体不是“失败就回滚重扫”", 1)
    add_text(doc, "一次申请由多个独立持久化步骤组成。某一步失败时，已经提交的前置步骤和 Checkpoint 不会全部回滚；重试只恢复当前 step。")
    add_flow(doc, [
        ("DOCUMENT_INTAKE 已提交", GREEN),
        ("↓", WHITE),
        ("DOCUMENT_ANALYSIS 已提交", GREEN),
        ("↓", WHITE),
        ("FRAUD_ASSESSMENT 失败", RED),
        ("↓ 延迟重试同一 eventId + 同一 step", WHITE),
        ("只重试 FRAUD_ASSESSMENT", YELLOW),
    ])
    add_text(doc, "Outbox 的状态也不会因为消费业务失败自动退回 PENDING。它已完成“把消息交给 Broker”的职责；后续恢复属于 RabbitMQ 消费重试。")

    add_heading(doc, "9. 关键配置与代码索引", 1)
    add_table(doc, ["主题", "文件与位置", "重点"], [
        ["MQ 拓扑", "config/ApprovalRabbitConfig.java:29-133", "主、重试、死信交换机与队列"],
        ["消息结构", "messaging/ApprovalStartMessage.java:6-14", "eventId、step、retryCount"],
        ["创建 Outbox", "service/impl/ApprovalTaskServiceImpl.java:43-85", "首步骤与确定性子 eventId"],
        ["发布 Outbox", "messaging/ApprovalOutboxPublisher.java:39-60", "扫描、claim、Confirm、发布重试"],
        ["消费主队列", "messaging/ApprovalMessageConsumer.java:85-150", "ACK/NACK、业务重试、DLQ"],
        ["消费幂等", "service/impl/ApprovalConsumeServiceImpl.java:31-70", "租约、claimToken、完成状态"],
        ["单步执行", "service/impl/ApprovalSingleStepExecutor.java:46-161", "步骤分发与 nextStep"],
        ["Outbox SQL", "resources/mapper/ApprovalOutboxMapper.xml", "状态条件更新与扫描"],
        ["消费日志 SQL", "resources/mapper/ApprovalMessageConsumeLogMapper.xml", "幂等、失败与租约接管"],
        ["执行锁 SQL", "resources/mapper/ApprovalExecutionLockMapper.xml", "applicationId 互斥"],
        ["运行参数", "resources/application.yml:23-38,83-90", "Confirm、ACK、并发、重试"],
    ], [1900, 4300, 3160], font_size=8.5)

    add_heading(doc, "10. 调试与验证指南", 1)
    add_heading(doc, "10.1 推荐断点", 2)
    for item in [
        "DocumentServiceImpl.uploadDocument()：确认材料齐全后是否进入 submit。",
        "ApprovalTaskServiceImpl.persistStep()：观察 eventId、step 和 Outbox PENDING。",
        "ApprovalOutboxPublisher.publishPending()：观察 claim 和 Publisher Confirm。",
        "ApprovalMessageConsumer.consume()：观察 retryCount、ACK/NACK 和异常分支。",
        "ApprovalConsumeServiceImpl.tryStart()：观察 claimToken 与 leaseUntil。",
        "ApprovalSingleStepExecutor.execute()：观察当前 step 和返回的 nextStep。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "10.2 推荐 SQL", 2)
    add_code(doc, "SELECT id,event_id,application_id,event_type,publish_status,retry_count,next_retry_at,last_error\nFROM approval_outbox ORDER BY id DESC;\n\nSELECT event_id,application_id,consume_status,retry_count,attempt_no,lease_until,last_error\nFROM approval_message_consume_log ORDER BY id DESC;\n\nSELECT * FROM approval_execution_lock;\nSELECT * FROM approval_step_checkpoint ORDER BY id DESC;\nSELECT * FROM manual_review_ticket ORDER BY ticket_id DESC;")
    add_heading(doc, "10.3 三个必做实验", 2)
    add_number(doc, "停止 RabbitMQ 后上传完整材料：确认 Outbox 进入 RETRY；恢复 RabbitMQ 后变为 PUBLISHED/CONSUMED。")
    add_number(doc, "让某审批步骤稳定抛异常：观察 retry.queue 等待 TTL、回主队列及 retryCount 增长。")
    add_number(doc, "持续失败超过上限：确认申请进入 MANUAL_REVIEW、生成工单，消息留在 approval.dead.queue。")
    add_heading(doc, "11. 当前实现的改进建议", 1)
    add_table(doc, ["优先级", "问题", "建议"], [
        ["P0", "重试/死信发布仅调用 convertAndSend，未等待 Confirm", "也使用 Publisher Confirm，或把重试与死信写入专用 Outbox"],
        ["P0", "NACK + requeue 无退避和上限，可能热循环", "配置容器级退避、设置异常计数，并在阈值后告警/隔离"],
        ["P1", "锁竞争无最大等待边界", "增加首次冲突时间、最长等待、独立计数及监控"],
        ["P1", "DLQ 当前没有补偿消费者", "建设管理命令：查看、校验状态、重新发布或关闭事件"],
        ["P1", "Publisher 逐条同步等待最多 10 秒", "使用异步 Confirm 或有界并行，提高 Broker 故障时吞吐"],
        ["P2", "需要验证不可路由消息", "配置 mandatory/ReturnsCallback，将 Return 视为发布失败"],
    ], [1100, 3700, 4560], font_size=8.8)
    add_callout(doc, "最终记忆", "Outbox 保证业务提交后“总有一条待发事件”；Publisher Confirm 尽量保证它到达 Broker；消费端用 eventId 幂等吸收重复；业务失败进入延迟重试，耗尽后转人工和 DLQ。", MID_BLUE)

    # Keep tables from splitting headers awkwardly and set metadata.
    props = doc.core_properties
    props.title = "FinancialCrisis 项目 MQ Outbox 链路与重试机制"
    props.subject = "RabbitMQ、Transactional Outbox、重试队列、死信队列代码导读"
    props.author = "Codex"
    props.keywords = "RabbitMQ, Outbox, 重试队列, 死信队列, 幂等消费"
    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    build()
